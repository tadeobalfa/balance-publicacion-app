
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
import streamlit as st
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from openpyxl import load_workbook
from openpyxl.utils.cell import range_boundaries


# ============================================================
# CONFIGURACIÓN DE PÁGINA
# ============================================================

st.set_page_config(
    page_title="Balance Publicación a Word",
    page_icon="📄",
    layout="wide",
)

st.title("📄 Generador de Word - Balance de Publicación")
st.caption(
    "Subí el Excel con áreas de impresión definidas. "
    "Cada área de impresión se convierte en una página del Word."
)


# ============================================================
# UTILIDADES
# ============================================================

def safe_filename(name: str) -> str:
    name = (name or "Balance Publicacion.docx").strip()
    name = re.sub(r'[\\/:*?"<>|]', "", name)
    if not name.lower().endswith(".docx"):
        name += ".docx"
    return name


def save_upload(uploaded_file) -> str:
    temp_dir = tempfile.mkdtemp(prefix="balance_publicacion_")
    path = os.path.join(temp_dir, uploaded_file.name)
    with open(path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return path


def find_soffice() -> str | None:
    candidates = [
        os.environ.get("SOFFICE_PATH"),
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in candidates:
        if c and os.path.exists(c):
            return c
    return None


def normalize_print_area(print_area) -> str:
    if print_area is None:
        return ""
    return str(print_area)


def split_print_areas(print_area: str) -> list[str]:
    """
    Divide un PrintArea de Excel con múltiples áreas.
    Ejemplo:
    'Hoja'!$A$1:$E$50,'Hoja'!$A$55:$E$100
    """
    if not print_area:
        return []

    parts = []
    current = []
    in_quote = False

    for ch in print_area:
        if ch == "'":
            in_quote = not in_quote
            current.append(ch)
        elif ch == "," and not in_quote:
            parts.append("".join(current).strip())
            current = []
        else:
            current.append(ch)

    if current:
        parts.append("".join(current).strip())

    ranges = []
    for part in parts:
        if "!" in part:
            part = part.split("!", 1)[1]
        part = part.replace("'", "").strip()
        if part:
            ranges.append(part)

    return ranges


def clean_range(area: str) -> str:
    area = area.strip().upper().replace("$", "")
    if "!" in area:
        area = area.split("!", 1)[1].replace("'", "")
    range_boundaries(area)  # valida
    return area


def range_size_score(ws, area: str) -> tuple[float, float]:
    """
    Estima ancho y alto visual usando ancho de columnas y alto de filas.
    """
    min_col, min_row, max_col, max_row = range_boundaries(area)

    width = 0.0
    for c in range(min_col, max_col + 1):
        letter = ws.cell(1, c).column_letter
        dim = ws.column_dimensions.get(letter)
        width += float((dim.width if dim is not None and dim.width is not None else 8.43))

    height = 0.0
    for r in range(min_row, max_row + 1):
        dim = ws.row_dimensions.get(r)
        height += float((dim.height if dim is not None and dim.height is not None else 15))

    return width, height


def auto_orientation(ws, area: str) -> str:
    """
    Devuelve:
    - V para vertical
    - H para horizontal
    """
    width, height = range_size_score(ws, area)

    # Relación práctica para detectar áreas anchas.
    return "H" if width * 7 > height * 0.85 else "V"


def make_area_lines(ws, areas: list[str]) -> str:
    lines = []
    for area in areas:
        area = clean_range(area)
        ori = auto_orientation(ws, area)
        lines.append(f"{area} | {ori}")
    return "\n".join(lines)


def parse_area_lines(text: str) -> list[tuple[str, str]]:
    result = []

    for line in text.splitlines():
        line = line.strip()

        if not line:
            continue

        if "|" in line:
            area, ori = [x.strip().upper() for x in line.split("|", 1)]
        else:
            area, ori = line.strip().upper(), "AUTO"

        area = clean_range(area)

        if ori not in ("V", "H", "AUTO"):
            raise ValueError(f"Orientación inválida en línea: {line}. Usá V, H o AUTO.")

        result.append((area, ori))

    return result


def get_print_areas_from_sheet(ws) -> list[str]:
    raw = normalize_print_area(ws.print_area)
    return [clean_range(a) for a in split_print_areas(raw)]


def infer_default_output_name(uploaded_name: str) -> str:
    name = Path(uploaded_name).stem
    name = re.sub(r"(?i)^BAL\s*PUBLIC\s*-?\s*", "", name).strip()
    if not name:
        name = "Balance Publicacion"
    return f"Bce. 2025 - {name}.docx"


# ============================================================
# RENDER DE EXCEL CON LIBREOFFICE
# ============================================================

def hide_all_except(wb, sheet_name: str):
    """
    Deja visible solo la hoja elegida.
    LibreOffice exporta hojas visibles, por eso ocultamos el resto.
    """
    for ws in wb.worksheets:
        ws.sheet_state = "visible" if ws.title == sheet_name else "hidden"
    wb.active = wb.sheetnames.index(sheet_name)


def prepare_single_area_workbook(
    original_xlsx: str,
    output_xlsx: str,
    sheet_name: str,
    area: str,
    orientation: str,
):
    """
    Crea una copia temporal con una sola área de impresión.
    """
    wb = load_workbook(original_xlsx)
    ws = wb[sheet_name]

    hide_all_except(wb, sheet_name)

    ws.print_area = area
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.orientation = "landscape" if orientation == "H" else "portrait"

    # Márgenes chicos.
    ws.page_margins.left = 0.25
    ws.page_margins.right = 0.25
    ws.page_margins.top = 0.25
    ws.page_margins.bottom = 0.25
    ws.page_margins.header = 0
    ws.page_margins.footer = 0

    wb.save(output_xlsx)


def libreoffice_to_pdf(soffice: str, xlsx_path: str, out_dir: str) -> str:
    os.makedirs(out_dir, exist_ok=True)

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to",
        "pdf",
        "--outdir",
        out_dir,
        xlsx_path,
    ]

    res = subprocess.run(cmd, capture_output=True, text=True, timeout=180)

    expected = os.path.join(out_dir, Path(xlsx_path).with_suffix(".pdf").name)

    if res.returncode != 0 or not os.path.exists(expected):
        raise RuntimeError(
            "LibreOffice no pudo convertir el Excel a PDF.\n\n"
            f"Comando: {' '.join(cmd)}\n\n"
            f"STDOUT:\n{res.stdout}\n\nSTDERR:\n{res.stderr}"
        )

    return expected


def pdf_first_page_to_png(pdf_path: str, png_path: str, zoom: float = 3.0):
    pdf = fitz.open(pdf_path)

    if len(pdf) == 0:
        pdf.close()
        raise RuntimeError(f"El PDF generado no tiene páginas: {pdf_path}")

    page = pdf[0]
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    pix.save(png_path)
    pdf.close()


# ============================================================
# WORD
# ============================================================

def add_image_page(doc: Document, png_path: str, orientation: str, first: bool):
    if not first:
        doc.add_section()

    section = doc.sections[-1]

    if orientation == "H":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # Márgenes 0 porque la imagen ya viene renderizada como página.
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    usable_width = section.page_width
    usable_height = section.page_height

    p = doc.add_paragraph()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0

    run = p.add_run()
    pic = run.add_picture(png_path, width=usable_width)

    if pic.height > usable_height:
        pic.height = usable_height


def build_word(
    original_xlsx: str,
    soffice: str,
    caratula_sheet: str,
    caratula_area: str,
    balance_sheet: str,
    balance_areas: list[tuple[str, str]],
    output_name: str,
):
    temp_dir = tempfile.mkdtemp(prefix="balance_render_")
    doc = Document()

    pages = []

    # Página 1: carátula, siempre vertical.
    pages.append((caratula_sheet, caratula_area, "V", "Carátula"))

    wb_read = load_workbook(original_xlsx, read_only=False, data_only=False)
    ws_bal_read = wb_read[balance_sheet]

    for idx, (area, ori) in enumerate(balance_areas, start=1):
        if ori == "AUTO":
            ori = auto_orientation(ws_bal_read, area)
        pages.append((balance_sheet, area, ori, f"Balance {idx}"))

    progress = st.progress(0)
    status = st.empty()

    for idx, (sheet, area, ori, label) in enumerate(pages, start=1):
        status.write(
            f"Procesando página {idx}/{len(pages)}: "
            f"{sheet} - {area} - {'Horizontal' if ori == 'H' else 'Vertical'}"
        )

        xlsx_tmp = os.path.join(temp_dir, f"page_{idx:03d}.xlsx")
        pdf_tmp_dir = os.path.join(temp_dir, f"pdf_{idx:03d}")
        png_tmp = os.path.join(temp_dir, f"page_{idx:03d}.png")

        prepare_single_area_workbook(original_xlsx, xlsx_tmp, sheet, area, ori)
        pdf_path = libreoffice_to_pdf(soffice, xlsx_tmp, pdf_tmp_dir)
        pdf_first_page_to_png(pdf_path, png_tmp)
        add_image_page(doc, png_tmp, ori, first=(idx == 1))

        progress.progress(idx / len(pages))

    output_path = os.path.join(temp_dir, safe_filename(output_name))
    doc.save(output_path)

    status.write("Listo.")
    return output_path, len(pages)


# ============================================================
# UI
# ============================================================

soffice_path = find_soffice()

with st.expander("Estado del servidor", expanded=False):
    if soffice_path:
        st.success(f"LibreOffice detectado en el servidor: {soffice_path}")
    else:
        st.error(
            "No se detectó LibreOffice en el servidor. "
            "En Render esto se instala desde el Dockerfile."
        )

uploaded = st.file_uploader(
    "Cargar Excel con áreas de impresión definidas",
    type=["xlsx", "xlsm"],
)

if uploaded:
    xlsx_path = save_upload(uploaded)

    try:
        wb = load_workbook(xlsx_path, read_only=False, data_only=False)
    except Exception as e:
        st.error(f"No pude abrir el Excel. Error: {e}")
        st.stop()

    sheetnames = wb.sheetnames

    st.success("Excel cargado correctamente.")

    # Sugerir hojas.
    car_idx = 0
    bal_idx = 0

    for i, name in enumerate(sheetnames):
        low = name.lower()
        if "car" in low:
            car_idx = i
        if "bal public" in low or "reex" in low:
            bal_idx = i

    col1, col2 = st.columns(2)

    with col1:
        car_sheet = st.selectbox("Hoja de carátula", sheetnames, index=car_idx)

    with col2:
        bal_sheet = st.selectbox("Hoja del balance", sheetnames, index=bal_idx)

    ws_car = wb[car_sheet]
    ws_bal = wb[bal_sheet]

    car_areas = get_print_areas_from_sheet(ws_car)
    bal_areas = get_print_areas_from_sheet(ws_bal)

    st.subheader("Áreas detectadas")

    if not car_areas:
        st.warning("La hoja de carátula no tiene área de impresión definida.")
        car_area_default = ""
    else:
        car_area_default = car_areas[0]

    if not bal_areas:
        st.warning("La hoja del balance no tiene áreas de impresión definidas.")
        balance_lines_default = ""
    else:
        balance_lines_default = make_area_lines(ws_bal, bal_areas)

    car_area = st.text_input("Área de impresión de carátula", value=car_area_default)

    st.write(
        "Áreas del balance. Formato: `RANGO | V` o `RANGO | H`. "
        "También podés usar `AUTO`."
    )

    balance_area_text = st.text_area(
        "Áreas de impresión del balance, una por línea",
        value=balance_lines_default,
        height=300,
    )

    try:
        parsed_areas = parse_area_lines(balance_area_text)
        st.info(
            f"Se generarán {1 + len(parsed_areas)} páginas: "
            f"1 carátula + {len(parsed_areas)} páginas del balance."
        )
    except Exception as e:
        parsed_areas = []
        st.error(f"Hay un error en las áreas: {e}")

    output_name = st.text_input(
        "Nombre del Word de salida",
        value=infer_default_output_name(uploaded.name),
    )

    st.markdown("---")

    can_generate = bool(soffice_path and car_area and parsed_areas)

    if st.button("Generar Word", type="primary", disabled=not can_generate):
        try:
            clean_range(car_area)

            with st.spinner("Generando Word. Esto puede tardar unos minutos..."):
                output_path, total_pages = build_word(
                    original_xlsx=xlsx_path,
                    soffice=soffice_path,
                    caratula_sheet=car_sheet,
                    caratula_area=clean_range(car_area),
                    balance_sheet=bal_sheet,
                    balance_areas=parsed_areas,
                    output_name=output_name,
                )

            st.success(f"Word generado correctamente. Total de páginas: {total_pages}")

            with open(output_path, "rb") as f:
                st.download_button(
                    label="⬇️ Descargar Word",
                    data=f,
                    file_name=safe_filename(output_name),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

        except Exception as e:
            st.error(f"Ocurrió un error: {e}")
            st.warning(
                "Revisá que el Excel tenga áreas de impresión definidas y que "
                "las hojas seleccionadas sean correctas."
            )

else:
    st.info(
        "Prepará el Excel definiendo un área de impresión para la carátula "
        "y varias áreas de impresión en la hoja del balance."
    )

